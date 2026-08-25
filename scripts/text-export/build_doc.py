# -*- coding: utf-8 -*-
"""Сборка выгрузки текстов со снимками разделов и графой для правок."""
import json, os, re, sys
from PIL import Image
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SP = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(SP, 'shots')
JPG = os.path.join(SP, 'jpg')
OUT = os.environ.get('OUT_DOCX', '/Users/buzanovsergey/Desktop/Тексты сайта ДПО – со снимками страниц – 25.08.2026.docx')
PREVIEW = os.environ.get('PREVIEW') == '1'   # укороченный файл для проверки вёрстки

GRAY = RGBColor(0x60, 0x5A, 0x52)
BLUE = RGBColor(0x16, 0x58, 0xDA)
CONTENT_CM = 17.0

doc_data = json.load(open(os.path.join(SP, 'doc.json')))
manifest = json.load(open(os.path.join(SHOTS, 'manifest.json')))

# ---------- снимки: PNG -> JPEG, чтобы файл не разнесло до десятка мегабайт ----------
os.makedirs(JPG, exist_ok=True)
def as_jpeg(name):
    dst = os.path.join(JPG, name.replace('.png', '.jpg'))
    if not os.path.exists(dst):
        im = Image.open(os.path.join(SHOTS, name)).convert('RGB')
        im.save(dst, 'JPEG', quality=90, optimize=True, subsampling=0)
    return dst

# ---------- разбор манифеста: страница -> название раздела -> очередь снимков ----------
shots_by_page = {}
for e in manifest:
    if e.get('error') or not e.get('slices'):
        continue
    base = e['section'].split('#')[0]
    shots_by_page.setdefault(e['page'], {}).setdefault(base, []).append(e)

def take_shot(page, section):
    q = shots_by_page.get(page, {}).get(section)
    return q.pop(0) if q else None

# ---------- документ ----------
d = Document()
s = d.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)          # A4
s.left_margin = s.right_margin = Cm(2.0)
s.top_margin = s.bottom_margin = Cm(1.5)

normal = d.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(14)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
normal.element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(6)

def para(text='', *, size=14, bold=False, italic=False, color=None,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, container=None):
    p = (container or d).add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    return p

def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:color'), 'auto'); el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)

# Порядок детей <w:tblPr> задан схемой; вставка не на своё место делает файл
# «нечитаемым содержимым» для Word, поэтому кладём элементы по позиции.
TBLPR_SEQ = ('w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual',
             'w:tblStyleRowBandSize', 'w:tblStyleColBandSize', 'w:tblW', 'w:jc',
             'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders', 'w:shd', 'w:tblLayout',
             'w:tblCellMar', 'w:tblLook', 'w:tblCaption', 'w:tblDescription', 'w:tblPrChange')


def tblpr_put(tblPr, tag):
    existing = tblPr.find(qn(tag))
    if existing is not None:
        return existing
    el = OxmlElement(tag)
    i = TBLPR_SEQ.index(tag)
    tblPr.insert_element_before(el, *TBLPR_SEQ[i + 1:])
    return el


def borders(table, color='BFB9B0', sz=4):
    b = tblpr_put(table._tbl.tblPr, 'w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)


def fixed_layout(table, total_cm, widths_cm):
    table.autofit = False                      # сам ставит <w:tblLayout type="fixed">
    w = tblpr_put(table._tbl.tblPr, 'w:tblW')
    w.set(qn('w:w'), str(int(total_cm * 567))); w.set(qn('w:type'), 'dxa')
    grid = table._tbl.find(qn('w:tblGrid'))
    for col, cm in zip(grid.findall(qn('w:gridCol')), widths_cm):
        col.set(qn('w:w'), str(int(cm * 567)))


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true')
    trPr.append(el)

def no_split(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:cantSplit'); el.set(qn('w:val'), 'true')
    trPr.append(el)

def cell_text(cell, text, *, size=14, italic=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if text:
        r = p.add_run(text)
        r.italic = italic
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    return p

def add_shot(entry, label):
    """Снимок раздела: одна или несколько полос. Снимок с телефона ставится узким –
    во всю ширину полосы он выглядел бы плакатом, а не экраном телефона."""
    pic_cm = CONTENT_CM if entry.get('width', 1440) >= 1440 else 8.0
    for i, sl in enumerate(entry['slices']):
        p = d.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(as_jpeg(sl['file']), width=Cm(pic_cm))
        if len(entry['slices']) > 1:
            cap = f'{label} – часть {i + 1} из {len(entry["slices"])}'
        else:
            cap = label
        para(cap, size=10, italic=True, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

def add_table(items):
    t = d.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    widths_cm = (3.2, 7.8, 6.0)
    widths = tuple(Cm(x) for x in widths_cm)
    borders(t)
    fixed_layout(t, sum(widths_cm), widths_cm)
    hdr = t.rows[0]
    repeat_header(hdr)
    for cell, title, w in zip(hdr.cells, ('Что это', 'Текст на сайте', 'Новый текст'), widths):
        cell.width = w
        shade(cell, 'EFECE6')
        p = cell_text(cell, title, size=12)
        p.runs[0].bold = True
    for it in items:
        row = t.add_row()
        no_split(row)
        c0, c1, c2 = row.cells
        c0.width, c1.width, c2.width = widths
        cell_text(c0, it['type'], size=11, italic=True, color=GRAY)
        cell_text(c1, it['text'], size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
        c2.width = widths[2]
        cell_text(c2, '', size=14)
    para(after=10)
    return t

def page_numbers(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.font.size = Pt(11); r.font.color.rgb = GRAY
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    r._r.addnext(fld)


page_numbers(s)

# ---------- титул и памятка ----------
if PREVIEW:
    pi = int(os.environ.get('PV_PAGE', '0'))
    a, b = (int(v) for v in os.environ.get('PV_SEC', '0:3').split(':'))
    doc_data['pages'] = [dict(doc_data['pages'][pi],
                              sections=doc_data['pages'][pi]['sections'][a:b])]
if not PREVIEW:
    para('Тексты сайта, которые правятся руками', size=24, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    para('Центр ДПО факультета права НИУ ВШЭ · снимки страниц и графа для правок',
         size=16, color=GRAY, align=WD_ALIGN_PARAGRAPH.LEFT, after=16)

    para('Как пользоваться файлом', size=15, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)
    for line in [
        'Файл повторяет сайт сверху вниз: страница – раздел – тексты этого раздела. '
        'Перед текстами каждого раздела стоит снимок соответствующего куска страницы, '
        'поэтому видно, где именно на сайте живёт каждая строка.',
        'Правки пишите в третью графу «Новый текст». Вторую графу («Текст на сайте») '
        'менять не нужно: она нужна, чтобы было с чем сверять. Пустая третья графа '
        'означает «оставить как есть».',
        'Если строку надо убрать с сайта совсем – напишите в третьей графе слово УДАЛИТЬ.',
        'Названия и описания программ в этот файл не входят: они тянутся из каталога '
        'НИУ ВШЭ и правятся там, а не на сайте. Они собраны в отдельном файле '
        '«Тексты сайта ДПО – из каталога».',
        'Девизы первого экрана меняются сами при каждом заходе, их список – в отдельном '
        'файле «Девизы первого экрана».',
    ]:
        p = para(line, size=14, after=6)
        p.paragraph_format.first_line_indent = Cm(0.75)

    para('Что означают пометки в графе «Что это»', size=15, bold=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=6)
    LEGEND = [
        ('Заголовок H1', 'главный заголовок страницы, он же виден поисковикам; на странице он один'),
        ('Заголовок H2, H3', 'заголовки разделов и карточек внутри страницы'),
        ('Надзаголовок', 'короткая строка над заголовком, обычно набрана капителью'),
        ('Текст', 'обычный текст, подпись, число или пояснение'),
        ('Ссылка', 'текст, по которому кликают и уходят на другую страницу'),
        ('Кнопка', 'надпись на кнопке'),
        ('Пункт списка', 'строка маркированного или нумерованного перечня'),
        ('Пункт выбора', 'вариант в выпадающем списке'),
    ]
    for name, expl in LEGEND:
        p = d.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f'[{name}] ')
        r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GRAY
        r2 = p.add_run('– ' + expl)
        r2.font.size = Pt(12)

    para('Снимки сделаны 25 августа 2026 года с рабочей версии сайта, ширина окна 1440 точек. '
         'Два снимка сделаны на телефоне – там, где элемент виден только на узком экране.',
         size=12, italic=True, color=GRAY, before=10, after=6)

# ---------- пояснения и уточнённые названия разделов ----------
# Ключ: (страница, название раздела, какой он по счёту среди одноимённых).
RENAME = {
    ('Каталог программ', 'Служебные элементы страницы', 0): 'Поиск и сортировка',
    ('Каталог программ', 'Служебные элементы страницы', 1): 'Панель внизу экрана на телефоне',
}
NOTES = {
    ('Главная', 'Первый экран', 0):
        'Девиз под заголовком меняется при каждом заходе на сайт – на снимке одна из '
        'восьми фраз. Их список правится в отдельном файле «Девизы первого экрана».',
    ('Главная', 'Ближайшие старты', 0):
        'Даты и названия программ в бегущей строке берутся из каталога НИУ ВШЭ и здесь не правятся.',
    ('Главная', 'Популярные программы', 0):
        'Названия, описания и цены карточек берутся из каталога НИУ ВШЭ. Здесь правятся только '
        'надзаголовок, заголовок, вводный текст и ссылка внизу.',
    ('Главная', 'Направления по сферам права', 0):
        'Перечни программ внутри карточек собираются из каталога. Здесь правятся заголовок, '
        'вводный текст и названия сфер.',
    ('Главная', 'Преподаватели', 0):
        'Имена, должности и фотографии подтягиваются с сайта НИУ ВШЭ.',
    ('Главная', 'Отзывы выпускников', 0):
        'Тексты отзывов лежат в каталоге программ, здесь правятся заголовок и вводный текст.',
    ('Каталог программ', 'Фильтры каталога', 0):
        'Числа в скобках считаются сами по каталогу и меняться руками не могут.',
    ('Каталог программ', 'Поиск и сортировка', 0):
        'Строка поиска, выпадающий список сортировки и счётчик найденного – под фильтрами.',
}

# ---------- страницы ----------
missing = []
for page in doc_data['pages']:
    ph = para(page['name'], size=20, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    ph.paragraph_format.page_break_before = True
    if page.get('note'):
        for line in page['note'].split('\n'):
            para(line, size=12, italic=True, color=GRAY,
                 align=WD_ALIGN_PARAGRAPH.LEFT, after=4)
    seen = {}
    for sec in page['sections']:
        nth = seen.get(sec['name'], 0)
        seen[sec['name']] = nth + 1
        shown = RENAME.get((page['name'], sec['name'], nth), sec['name'])
        if shown:
            para(shown, size=15, bold=True,
                 align=WD_ALIGN_PARAGRAPH.LEFT, before=14, after=4)
        entry = take_shot(page['name'], sec['name'])
        if entry:
            label = f'{page["name"]} · {shown}' if shown else page['name']
            add_shot(entry, 'Снимок: ' + label)
        else:
            missing.append((page['name'], sec['name']))
        note = NOTES.get((page['name'], shown, nth))
        if note:
            para(note, size=12, italic=True, color=GRAY, after=6)
        if sec.get('note'):
            for line in sec['note'].split('\n'):
                para(line, size=12, italic=True, color=GRAY, after=4)
        add_table(sec['items'])

d.save(OUT)
n_items = sum(len(s['items']) for pg in doc_data['pages'] for s in pg['sections'])
print('сохранено:', OUT)
print('размер: %.1f МБ' % (os.path.getsize(OUT) / 1e6))
print('страниц сайта:', len(doc_data['pages']), '· разделов:',
      sum(len(pg['sections']) for pg in doc_data['pages']), '· строк текста:', n_items)
if missing:
    print('БЕЗ СНИМКА:', missing)
